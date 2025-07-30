#!/usr/bin/env python3
"""
Script to create a test Word document with simulated comments for debugging
"""

from docx import Document
from docx.shared import RGBColor
import os

def create_test_document_with_markup():
    """Create a document with text that simulates comment-style changes"""
    
    doc = Document()
    
    # Add a title
    title = doc.add_heading('Test Document for Comment Analysis', 0)
    
    # Add some content with "comment-style" markup
    doc.add_paragraph('This is the original document with change instructions.')
    
    # Simulate comments using bracket notation
    para1 = doc.add_paragraph()
    para1.add_run('Johnny likes to play basketball. ')
    para1.add_run('[COMMENT: change Johnny to Jimmy everywhere]').font.color.rgb = RGBColor(255, 0, 0)
    
    para2 = doc.add_paragraph()
    para2.add_run('The weather is ')
    para2.add_run('nice').underline = True
    para2.add_run(' today. ')
    para2.add_run('[COMMENT: change nice to excellent]').font.color.rgb = RGBColor(255, 0, 0)
    
    para3 = doc.add_paragraph()
    para3.add_run('This sentence should be removed. ')
    para3.add_run('[COMMENT: delete this sentence]').font.color.rgb = RGBColor(255, 0, 0)
    
    para4 = doc.add_paragraph()
    para4.add_run('The project will be completed soon. ')
    para4.add_run('[COMMENT: add "very" before soon]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Save the original document
    original_path = 'test_original_with_comments.docx'
    doc.save(original_path)
    print(f'✅ Created: {original_path}')
    
    # Create a revised version
    doc_revised = Document()
    
    # Add the same title
    doc_revised.add_heading('Test Document for Comment Analysis', 0)
    doc_revised.add_paragraph('This is the original document with change instructions.')
    
    # Apply the changes
    doc_revised.add_paragraph('Jimmy likes to play basketball.')  # Johnny -> Jimmy
    doc_revised.add_paragraph('The weather is excellent today.')  # nice -> excellent
    # Skip the sentence that should be deleted
    doc_revised.add_paragraph('The project will be completed very soon.')  # added "very"
    
    # Save the revised document
    revised_path = 'test_revised_applied.docx'
    doc_revised.save(revised_path)
    print(f'✅ Created: {revised_path}')
    
    return original_path, revised_path

if __name__ == "__main__":
    original, revised = create_test_document_with_markup()
    print(f"\nTest files created:")
    print(f"  Original: {original}")
    print(f"  Revised:  {revised}")
    print(f"\nYou can now test the app with these files!")