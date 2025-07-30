#!/usr/bin/env python3
"""
Create realistic test documents with common real-world comment patterns
"""

from docx import Document
from docx.shared import RGBColor
import os

def create_realistic_documents():
    """Create documents with realistic comment patterns"""
    
    # Create original document with realistic comments
    doc = Document()
    
    # Add a title
    title = doc.add_heading('Project Report Draft', 0)
    
    # Add realistic content with various types of comments
    doc.add_paragraph('This report summarizes our findings from the recent market research project.')
    
    # Spelling correction
    para1 = doc.add_paragraph()
    para1.add_run('The ')
    para1.add_run('recieve').underline = True  # Misspelled word
    para1.add_run(' of feedback was overwhelming. ')
    para1.add_run('[COMMENT: correct spelling: receive]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Word choice correction  
    para2 = doc.add_paragraph()
    para2.add_run('Our team worked ')
    para2.add_run('good').underline = True
    para2.add_run(' together to complete the analysis. ')
    para2.add_run('[COMMENT: should be "well" not "good"]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Simple replacement
    para3 = doc.add_paragraph()
    para3.add_run('The results were ')
    para3.add_run('nice').underline = True
    para3.add_run(' and exceeded our expectations. ')
    para3.add_run('[COMMENT: excellent]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Global replacement
    para4 = doc.add_paragraph()
    para4.add_run('Company ABC provided valuable insights. ABC has been a reliable partner. ')
    para4.add_run('We recommend continuing our relationship with ABC. ')
    para4.add_run('[COMMENT: change all "ABC" to "Acme Corp"]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Deletion
    para5 = doc.add_paragraph()
    para5.add_run('The weather was terrible today, but ')
    para5.add_run('obviously ').underline = True
    para5.add_run('the project continued as planned. ')
    para5.add_run('[COMMENT: delete "obviously"]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Addition
    para6 = doc.add_paragraph()
    para6.add_run('The project will be completed soon. ')
    para6.add_run('[COMMENT: add "very" before "soon"]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Save the original document
    original_path = 'realistic_original_with_comments.docx'
    doc.save(original_path)
    print(f'✅ Created: {original_path}')
    
    # Create revised version with corrections applied
    doc_revised = Document()
    
    # Add the same title
    doc_revised.add_heading('Project Report Draft', 0)
    doc_revised.add_paragraph('This report summarizes our findings from the recent market research project.')
    
    # Apply corrections
    doc_revised.add_paragraph('The receive of feedback was overwhelming.')  # Fixed spelling
    doc_revised.add_paragraph('Our team worked well together to complete the analysis.')  # Fixed grammar
    doc_revised.add_paragraph('The results were excellent and exceeded our expectations.')  # Better word choice
    doc_revised.add_paragraph('Company Acme Corp provided valuable insights. Acme Corp has been a reliable partner. We recommend continuing our relationship with Acme Corp.')  # Global replacement
    doc_revised.add_paragraph('The weather was terrible today, but the project continued as planned.')  # Removed "obviously"
    doc_revised.add_paragraph('The project will be completed very soon.')  # Added "very"
    
    # Save the revised document
    revised_path = 'realistic_revised_applied.docx'
    doc_revised.save(revised_path)
    print(f'✅ Created: {revised_path}')
    
    return original_path, revised_path

if __name__ == "__main__":
    original, revised = create_realistic_documents()
    print(f"\nRealistic test files created:")
    print(f"  Original: {original}")
    print(f"  Revised:  {revised}")
    print(f"\nThese documents test:")
    print(f"  - Spelling corrections (recieve -> receive)")
    print(f"  - Grammar fixes (good -> well)")
    print(f"  - Word improvements (nice -> excellent)")
    print(f"  - Global replacements (ABC -> Acme Corp)")
    print(f"  - Deletions (remove 'obviously')")
    print(f"  - Additions (add 'very')")
    print(f"\nYou can now test the app with these realistic documents!")