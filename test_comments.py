#!/usr/bin/env python3
"""
Test script to debug Word comment extraction
"""

import sys
from docx import Document
import zipfile
import xml.etree.ElementTree as ET

def analyze_word_document(file_path):
    """Analyze a Word document to understand its structure"""
    print(f"🔍 Analyzing: {file_path}")
    
    try:
        # Method 1: Using python-docx
        print("\n📄 Method 1: Using python-docx")
        doc = Document(file_path)
        print(f"   Paragraphs: {len(doc.paragraphs)}")
        print(f"   Tables: {len(doc.tables)}")
        
        # Check for comments using python-docx
        if hasattr(doc.part, 'rels'):
            print(f"   Document relationships: {len(doc.part.rels)}")
            
            for rel_id, rel in doc.part.rels.items():
                print(f"   - {rel_id}: {rel.target_ref}")
        
        # Method 2: Direct ZIP analysis
        print("\n📦 Method 2: Direct ZIP file analysis")
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            file_list = docx_zip.namelist()
            print(f"   Total files in ZIP: {len(file_list)}")
            
            # Look for comments files
            comment_files = [f for f in file_list if 'comment' in f.lower()]
            print(f"   Comment-related files: {comment_files}")
            
            # Check for relationships
            rel_files = [f for f in file_list if 'rel' in f.lower()]
            print(f"   Relationship files: {rel_files}")
            
            # Try to read comments.xml if it exists
            if 'word/comments.xml' in file_list:
                print("\n💬 Found comments.xml!")
                comments_xml = docx_zip.read('word/comments.xml')
                
                # Parse the XML
                try:
                    root = ET.fromstring(comments_xml)
                    print(f"   XML root tag: {root.tag}")
                    
                    # Define namespace
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    # Find comments
                    comments = root.findall('.//w:comment', ns)
                    print(f"   Found {len(comments)} comment elements")
                    
                    for i, comment in enumerate(comments):
                        comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
                        
                        # Get comment text
                        text_elements = comment.findall('.//w:t', ns)
                        comment_text = ''.join(elem.text or '' for elem in text_elements)
                        
                        print(f"   Comment {i+1}: ID={comment_id}, Author={author}")
                        print(f"   Text: {comment_text[:100]}{'...' if len(comment_text) > 100 else ''}")
                        print()
                        
                except Exception as e:
                    print(f"   Error parsing comments XML: {e}")
            else:
                print("\n❌ No comments.xml found")
                
            # Check document.xml for comment references
            if 'word/document.xml' in file_list:
                print("\n📄 Checking document.xml for comment references")
                doc_xml = docx_zip.read('word/document.xml').decode('utf-8', errors='ignore')
                
                import re
                comment_refs = re.findall(r'commentReference.*?w:id="(\d+)"', doc_xml)
                print(f"   Found {len(comment_refs)} comment references: {comment_refs}")
                
        # Method 3: Check relationships
        print("\n🔗 Method 3: Checking relationships")
        try:
            doc_part = doc.part
            print(f"   Document part: {doc_part}")
            
            if hasattr(doc_part, 'rels'):
                print(f"   Relationships: {len(doc_part.rels)}")
                for rel_id, rel in doc_part.rels.items():
                    print(f"   - {rel_id}: {rel.target_ref}")
                    if 'comment' in rel.target_ref.lower():
                        print(f"     ✅ FOUND COMMENTS RELATIONSHIP!")
                        
        except Exception as e:
            print(f"   Error checking relationships: {e}")
            
    except Exception as e:
        print(f"❌ Error analyzing document: {e}")

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python test_comments.py <path_to_word_document.docx>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    analyze_word_document(file_path)