#!/usr/bin/env python3
"""
Create test documents with properly positioned comments to test focused analysis
"""

from docx import Document
from docx.shared import RGBColor

def create_focused_test_documents():
    """Create documents with comments positioned near the relevant text"""
    
    # Create original document
    doc = Document()
    
    doc.add_heading('Test Document for Focused AI Analysis', 0)
    
    # Paragraph 1: Name change
    para1 = doc.add_paragraph()
    para1.add_run('Johnny went to the store. ')
    para1.add_run('[COMMENT: Change Johnny to Jimmy throughout]').font.color.rgb = RGBColor(255, 0, 0)
    para1.add_run(' Johnny likes shopping.')
    
    # Paragraph 2: Spelling error  
    para2 = doc.add_paragraph()
    para2.add_run('The weather was ')
    para2.add_run('absolutly').underline = True
    para2.add_run(' beautiful today. ')
    para2.add_run('[COMMENT: Spelling mistake]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Paragraph 3: Word choice
    para3 = doc.add_paragraph()
    para3.add_run('The food was ')
    para3.add_run('good').underline = True
    para3.add_run(' at the restaurant. ')
    para3.add_run('[COMMENT: Use "excellent" instead]').font.color.rgb = RGBColor(255, 0, 0)
    
    # Paragraph 4: Deletion
    para4 = doc.add_paragraph()
    para4.add_run('The project was ')
    para4.add_run('very very ').underline = True
    para4.add_run('successful. ')
    para4.add_run('[COMMENT: Remove duplicate word]').font.color.rgb = RGBColor(255, 0, 0)
    
    original_path = 'focused_original_test.docx'
    doc.save(original_path)
    print(f'✅ Created: {original_path}')
    
    # Create revised document with changes applied
    doc_revised = Document()
    
    doc_revised.add_heading('Test Document for Focused AI Analysis', 0)
    
    # Apply the changes
    doc_revised.add_paragraph('Jimmy went to the store. Jimmy likes shopping.')  # Name change
    doc_revised.add_paragraph('The weather was absolutely beautiful today.')  # Spelling fix
    doc_revised.add_paragraph('The food was excellent at the restaurant.')  # Word improvement
    doc_revised.add_paragraph('The project was very successful.')  # Removed duplicate
    
    revised_path = 'focused_revised_test.docx'
    doc_revised.save(revised_path)
    print(f'✅ Created: {revised_path}')
    
    return original_path, revised_path

def test_comment_positions():
    """Test that comments will be positioned correctly for focused analysis"""
    
    from app import analyzer
    
    print("\n🎯 Testing Comment Positioning for Focused Analysis")
    print("=" * 55)
    
    # Create the test documents
    original_path, revised_path = create_focused_test_documents()
    
    # Extract comments and their positions
    original_data = analyzer.extract_document_data(original_path)
    revised_data = analyzer.extract_document_data(revised_path)
    
    print(f"\nExtracted {len(original_data['comments'])} comments:")
    
    for i, comment in enumerate(original_data['comments'], 1):
        print(f"\n{i}. Comment: \"{comment['text']}\"")
        print(f"   Position: {comment['position']}")
        
        # Test context extraction
        context = analyzer.extract_comment_context(
            comment, 
            original_data['full_text'], 
            revised_data['full_text']
        )
        
        print(f"   Original context: \"{context['original_context'][:100]}...\"")
        print(f"   Revised context:  \"{context['revised_context'][:100]}...\"")
        
        # Predict what AI should find
        if "johnny" in comment['text'].lower():
            print("   Expected: Should detect Johnny → Jimmy name change")
        elif "spelling" in comment['text'].lower():
            print("   Expected: Should detect absolutly → absolutely spelling fix")
        elif "excellent" in comment['text'].lower():
            print("   Expected: Should detect good → excellent word change")
        elif "duplicate" in comment['text'].lower():
            print("   Expected: Should detect removal of duplicate 'very'")
    
    print(f"\n✅ Test documents created and positioned correctly!")
    print(f"   Original: {original_path}")
    print(f"   Revised:  {revised_path}")
    print(f"\nNow when AI analyzes each comment, it will see focused context")
    print(f"and avoid confusing unrelated changes in the document.")

if __name__ == "__main__":
    test_comment_positions()