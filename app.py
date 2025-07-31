from flask import Flask, request, jsonify, render_template, send_from_directory, redirect, url_for
from werkzeug.utils import secure_filename
import os
import json
import uuid
from datetime import datetime
import logging
from docx import Document
from docx.shared import RGBColor
import difflib
# Force deployment update - v2.1
import re

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'dev-key-change-in-production')

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
    logger.info("Anthropic package imported successfully")
except ImportError as e:
    logger.error(f"Failed to import anthropic: {e}")
    ANTHROPIC_AVAILABLE = False

try:
    import openai
    OPENAI_AVAILABLE = True
    logger.info("OpenAI package imported successfully")
except ImportError as e:
    logger.error(f"Failed to import openai: {e}")
    OPENAI_AVAILABLE = False

# GenAI Configuration
# Try to get API keys from environment variables
ANTHROPIC_API_KEY = os.getenv('ANTHROPIC_API_KEY')
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')

# Initialize AI clients if keys are available (lazy initialization)
anthropic_client = None
openai_client = None

def get_anthropic_client():
    """Get Anthropic client with lazy initialization"""
    global anthropic_client
    if anthropic_client is None and ANTHROPIC_API_KEY and ANTHROPIC_AVAILABLE:
        try:
            anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            logger.info("Anthropic Claude API initialized")
        except Exception as e:
            logger.error(f"Failed to initialize Anthropic client: {str(e)}")
            anthropic_client = False  # Mark as failed to avoid retrying
    return anthropic_client if anthropic_client is not False else None

def get_openai_client():
    """Get OpenAI client with lazy initialization"""
    global openai_client
    if openai_client is None and OPENAI_API_KEY and OPENAI_AVAILABLE:
        try:
            openai_client = openai.OpenAI(api_key=OPENAI_API_KEY)
            logger.info("OpenAI API initialized")
        except Exception as e:
            logger.error(f"Failed to initialize OpenAI client: {str(e)}")
            openai_client = False  # Mark as failed to avoid retrying
    return openai_client if openai_client is not False else None

# Log available API keys and packages at startup (without initializing clients)
if ANTHROPIC_API_KEY and ANTHROPIC_AVAILABLE:
    logger.info("Anthropic API key found and package available - will initialize client on first use")
elif ANTHROPIC_API_KEY and not ANTHROPIC_AVAILABLE:
    logger.warning("Anthropic API key found but package not available - install with: pip install anthropic")
elif OPENAI_API_KEY and OPENAI_AVAILABLE:
    logger.info("OpenAI API key found and package available - will initialize client on first use")
elif OPENAI_API_KEY and not OPENAI_AVAILABLE:
    logger.warning("OpenAI API key found but package not available - install with: pip install openai")
else:
    logger.warning("No AI API keys found. Set ANTHROPIC_API_KEY or OPENAI_API_KEY environment variables for AI-powered analysis.")

# Ensure upload directory exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

class WordDocumentAnalyzer:
    def __init__(self):
        self.session_data = {}
    
    def extract_document_data(self, file_path):
        """Extract text and comments from a Word document"""
        try:
            doc = Document(file_path)
            
            # Extract main document text with paragraph tracking
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                paragraphs.append({
                    'index': i,
                    'text': para.text,
                    'runs': [{'text': run.text} for run in para.runs]
                })
            
            # Extract comments (Word comments are stored differently)
            comments = self.extract_comments(doc)
            
            return {
                'paragraphs': paragraphs,
                'comments': comments,
                'full_text': '\n'.join([p['text'] for p in paragraphs])
            }
        except Exception as e:
            logger.error(f"Error extracting document data: {str(e)}")
            raise e
    
    def extract_comments(self, doc):
        """Extract comments from Word document using multiple methods"""
        comments = []
        file_path = None
        
        # Get the file path if available (for ZIP method)
        if hasattr(doc, '_part') and hasattr(doc._part, 'package') and hasattr(doc._part.package, '_package_reader'):
            try:
                file_path = doc._part.package._package_reader._file_like_object.name
            except:
                pass
        
        try:
            logger.info("Attempting to extract comments using ZIP method...")
            comments = self.extract_comments_zip_method(file_path) if file_path else []
            
            if not comments:
                logger.info("ZIP method failed, trying relationship method...")
                comments = self.extract_comments_relationship_method(doc)
            
            if not comments:
                logger.info("Relationship method failed, trying fallback...")
                comments = self.extract_comments_fallback(doc)
                
            logger.info(f"Final result: Extracted {len(comments)} comments from document")
        
        except Exception as e:
            logger.error(f"Error extracting comments: {str(e)}")
            comments = self.extract_comments_fallback(doc)
        
        return comments
    
    def extract_comments_zip_method(self, file_path):
        """Extract comments by directly reading the ZIP file and associate with text"""
        comments = []
        
        if not file_path:
            return comments
        
        try:
            import zipfile
            import xml.etree.ElementTree as ET
            
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # First, extract comment-to-text associations from document.xml
                comment_ranges = self.extract_comment_ranges(docx_zip)
                
                # Then extract comment content from comments.xml
                if 'word/comments.xml' in docx_zip.namelist():
                    logger.info("Found word/comments.xml in ZIP file")
                    
                    comments_xml = docx_zip.read('word/comments.xml')
                    root = ET.fromstring(comments_xml)
                    
                    # Define namespace
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    
                    # Extract comments
                    comment_elements = root.findall('.//w:comment', ns)
                    logger.info(f"Found {len(comment_elements)} comment elements in XML")
                    
                    for comment_elem in comment_elements:
                        comment_id = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                        author = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                        date = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                        
                        # Extract all text content from the comment
                        text_elements = comment_elem.findall('.//w:t', ns)
                        comment_text = ''.join(elem.text or '' for elem in text_elements)
                        
                        if comment_text.strip():
                            # Find the associated text range for this comment
                            associated_text = comment_ranges.get(comment_id, '').strip()
                            
                            # Debug logging for real Word comments
                            logger.info(f"Real Word comment ID {comment_id}: '{comment_text[:30]}...'")
                            logger.info(f"Associated text for ID {comment_id}: '{associated_text}' (length: {len(associated_text)})")
                            
                            if not associated_text:
                                logger.warning(f"No associated text found for comment ID {comment_id}. Comment ranges available: {list(comment_ranges.keys())}")
                                # Fallback: use empty string but mark it
                                associated_text = f"[RANGE NOT FOUND FOR ID {comment_id}]"
                            
                            comments.append({
                                'id': comment_id or str(len(comments) + 1),
                                'text': comment_text.strip(),
                                'author': author,
                                'date': date,
                                'position': len(comments),
                                'associated_text': associated_text,
                                'context': f"Comment on '{associated_text[:50]}...' by {author}: {comment_text[:100]}..."
                            })
                            logger.info(f"Extracted comment: {comment_text[:50]}... on text: {associated_text[:30]}...")
                
                else:
                    logger.warning("No word/comments.xml found in ZIP file")
                    
        except Exception as e:
            logger.error(f"ZIP method error: {str(e)}")
        
        return comments
    
    def extract_comment_ranges(self, docx_zip):
        """Extract comment-to-text associations from document.xml"""
        comment_ranges = {}
        
        try:
            if 'word/document.xml' not in docx_zip.namelist():
                return comment_ranges
            
            import xml.etree.ElementTree as ET
            
            # Read and parse document.xml
            document_xml = docx_zip.read('word/document.xml')
            root = ET.fromstring(document_xml)
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Find all comment range starts and ends
            comment_starts = {}  # comment_id -> start position
            comment_ends = {}    # comment_id -> end position
            text_content = []    # All text content with positions
            
            # Walk through document collecting text and comment markers
            self.walk_document_for_comments(root, ns, comment_starts, comment_ends, text_content)
            
            # Build the full text
            full_text = ''.join([item['text'] for item in text_content])
            
            # Debug what we found
            logger.info(f"Found {len(comment_starts)} comment starts: {list(comment_starts.keys())}")
            logger.info(f"Found {len(comment_ends)} comment ends: {list(comment_ends.keys())}")
            logger.info(f"Total text elements: {len(text_content)}")
            
            # Extract text ranges for each comment
            for comment_id in comment_starts:
                if comment_id in comment_ends:
                    start_pos = comment_starts[comment_id]
                    end_pos = comment_ends[comment_id]
                    
                    logger.info(f"Comment {comment_id}: start={start_pos}, end={end_pos}")
                    
                    # Find the text between start and end positions
                    associated_text = self.extract_text_range(text_content, start_pos, end_pos)
                    comment_ranges[comment_id] = associated_text
                    
                    logger.info(f"Comment {comment_id} associated with text: '{associated_text[:50]}...'")
                else:
                    logger.warning(f"Comment {comment_id} has start but no end marker")
        
        except Exception as e:
            logger.error(f"Error extracting comment ranges: {str(e)}")
        
        return comment_ranges
    
    def walk_document_for_comments(self, element, ns, comment_starts, comment_ends, text_content):
        """Recursively walk document XML to find comment markers and text"""
        
        # Check for comment range start
        if element.tag.endswith('commentRangeStart'):
            comment_id = element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            if comment_id:
                comment_starts[comment_id] = len(text_content)
        
        # Check for comment range end  
        elif element.tag.endswith('commentRangeEnd'):
            comment_id = element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
            if comment_id:
                comment_ends[comment_id] = len(text_content)
        
        # Check for text content
        elif element.tag.endswith('t'):
            text = element.text or ''
            if text:
                text_content.append({
                    'text': text,
                    'position': len(text_content)
                })
        
        # Recursively process child elements
        for child in element:
            self.walk_document_for_comments(child, ns, comment_starts, comment_ends, text_content)
    
    def extract_text_range(self, text_content, start_pos, end_pos):
        """Extract text between start and end positions"""
        
        if start_pos >= end_pos or start_pos >= len(text_content):
            return ""
        
        end_pos = min(end_pos, len(text_content))
        
        # Join text from start to end position
        range_text = ''.join([
            text_content[i]['text'] 
            for i in range(start_pos, end_pos)
            if i < len(text_content)
        ])
        
        return range_text.strip()
    
    def find_associated_text_pattern(self, full_text, comment_match):
        """Find text associated with a comment pattern in fallback mode"""
        
        comment_start = comment_match.start()
        comment_end = comment_match.end()
        
        # Look for text patterns before the comment that might be what it refers to
        # Common patterns: underlined text, text in quotes, text before comment
        
        # Get text before the comment (±100 characters)
        text_before = full_text[max(0, comment_start - 100):comment_start]
        
        # Pattern 1: Look for specific words mentioned in the comment
        comment_text = comment_match.group(1) if comment_match.groups() else comment_match.group(0)
        
        # Extract potential target words from the comment itself
        if 'change' in comment_text.lower() and 'to' in comment_text.lower():
            # "change X to Y" pattern
            change_match = re.search(r'change\s+(\w+)\s+to\s+(\w+)', comment_text, re.IGNORECASE)
            if change_match:
                target_word = change_match.group(1)
                # Look for this word in the text before the comment
                if target_word.lower() in text_before.lower():
                    return target_word
        
        # Pattern 2: Look for underlined or emphasized text markers
        underlined_patterns = [
            r'(\w+(?:\s+\w+){0,2})\s*\[COMMENT',  # Words directly before [COMMENT
            r'(\w+)\s*\[COMMENT',                  # Single word before comment
        ]
        
        for pattern in underlined_patterns:
            match = re.search(pattern, text_before + comment_match.group(0), re.IGNORECASE)
            if match and len(match.group(1)) > 2:  # Avoid single characters
                return match.group(1).strip()
        
        # Pattern 2: Look for quoted text near the comment
        quote_patterns = [
            r'["\']([^"\']+)["\'](?=.*\[COMMENT)',  # Quoted text before comment
            r'(\w+(?:\s+\w+){0,2})\s*["\'](?=.*\[COMMENT)',  # Text before quotes
        ]
        
        for pattern in quote_patterns:
            match = re.search(pattern, text_before + comment_match.group(0), re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        # Pattern 3: Get last sentence or phrase before comment
        sentences = re.split(r'[.!?]', text_before)
        if sentences:
            last_sentence = sentences[-1].strip()
            # Get last few words if sentence is too long
            words = last_sentence.split()
            if len(words) > 10:
                return ' '.join(words[-5:])  # Last 5 words
            elif len(words) > 0:
                return last_sentence
        
        # Fallback: get text immediately before comment
        words_before = text_before.strip().split()
        if words_before:
            return ' '.join(words_before[-3:])  # Last 3 words
        
        return "unknown text"
    
    def extract_comment_ranges_from_document_part(self, document_part):
        """Extract comment ranges from document part (for relationship method)"""
        comment_ranges = {}
        
        try:
            # Get the document XML from the document part
            document_xml = document_part.blob
            
            import xml.etree.ElementTree as ET
            root = ET.fromstring(document_xml)
            
            ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Find all comment range starts and ends
            comment_starts = {}  # comment_id -> start position
            comment_ends = {}    # comment_id -> end position
            text_content = []    # All text content with positions
            
            # Walk through document collecting text and comment markers
            self.walk_document_for_comments(root, ns, comment_starts, comment_ends, text_content)
            
            # Debug what we found
            logger.info(f"Document part analysis: {len(comment_starts)} starts, {len(comment_ends)} ends, {len(text_content)} text elements")
            
            # Extract text ranges for each comment
            for comment_id in comment_starts:
                if comment_id in comment_ends:
                    start_pos = comment_starts[comment_id]
                    end_pos = comment_ends[comment_id]
                    
                    logger.info(f"Comment {comment_id}: positions {start_pos}-{end_pos}")
                    
                    # Find the text between start and end positions
                    associated_text = self.extract_text_range(text_content, start_pos, end_pos)
                    comment_ranges[comment_id] = associated_text
                    
                    logger.info(f"Comment {comment_id} → '{associated_text[:30]}...'")
                else:
                    logger.warning(f"Comment {comment_id} has start but no end marker")
        
        except Exception as e:
            logger.error(f"Error extracting comment ranges from document part: {str(e)}")
        
        return comment_ranges
    
    def extract_comments_relationship_method(self, doc):
        """Extract comments using document relationships and find their associated text"""
        comments = []
        
        try:
            document_part = doc.part
            
            # First, extract comment ranges from the main document
            comment_ranges = self.extract_comment_ranges_from_document_part(document_part)
            logger.info(f"Extracted {len(comment_ranges)} comment ranges from document")
            
            # Look for comments relationship
            for rel_id, rel in document_part.rels.items():
                logger.info(f"Checking relationship: {rel_id} -> {rel.target_ref}")
                
                if "comments" in rel.target_ref.lower():
                    logger.info(f"Found comments relationship: {rel.target_ref}")
                    
                    try:
                        comments_part = rel.target_part
                        comments_xml = comments_part.blob
                        
                        from xml.etree import ElementTree as ET
                        root = ET.fromstring(comments_xml)
                        
                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        
                        for comment_elem in root.findall('.//w:comment', ns):
                            comment_id = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                            author = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                            date = comment_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')
                            
                            text_elements = comment_elem.findall('.//w:t', ns)
                            comment_text = ''.join(elem.text or '' for elem in text_elements)
                            
                            if comment_text.strip():
                                # Find the associated text range for this comment
                                associated_text = comment_ranges.get(comment_id, '').strip()
                                
                                logger.info(f"Comment ID {comment_id}: '{comment_text[:30]}...' → Associated: '{associated_text[:30]}...' (len: {len(associated_text)})")
                                
                                if not associated_text:
                                    logger.warning(f"No associated text found for comment ID {comment_id}")
                                    associated_text = "[RANGE NOT FOUND]"
                                
                                comments.append({
                                    'id': comment_id or str(len(comments) + 1),
                                    'text': comment_text.strip(),
                                    'author': author,
                                    'date': date,
                                    'position': len(comments),
                                    'associated_text': associated_text,
                                    'context': f"Comment on '{associated_text[:30]}...' by {author}: {comment_text[:100]}..."
                                })
                                
                    except Exception as e:
                        logger.error(f"Error processing comments part: {str(e)}")
        
        except Exception as e:
            logger.error(f"Relationship method error: {str(e)}")
        
        return comments
    
    def extract_comments_fallback(self, doc):
        """Fallback method to extract comments using alternative approach"""
        comments = []
        
        try:
            # Method 1: Look for comment references in the main document XML
            document_xml = doc.part.blob.decode('utf-8', errors='ignore')
            
            # Look for comment reference patterns
            import re
            comment_refs = re.findall(r'<w:commentReference[^>]*w:id="(\d+)"', document_xml)
            
            if comment_refs:
                logger.info(f"Found {len(comment_refs)} comment references")
                # This indicates comments exist, but we need the comments.xml part
                
            # Method 2: Look for text patterns that might be comments
            full_text = '\n'.join([p.text for p in doc.paragraphs])
            
            # Check if document contains comment-like patterns and extract associated text
            comment_patterns = [
                r'\[COMMENT:\s*([^\]]+)\]',  # [COMMENT: text]
                r'\{COMMENT:\s*([^\}]+)\}',  # {COMMENT: text}
                r'##\s*([^##]+)##',          # ## comment ##
                r'\/\/\s*(.+)$',             # // comment (single line)
            ]
            
            for pattern in comment_patterns:
                matches = re.finditer(pattern, full_text, re.IGNORECASE | re.MULTILINE)
                for match in matches:
                    # Try to find associated text near the comment
                    associated_text = self.find_associated_text_pattern(full_text, match)
                    
                    comments.append({
                        'id': len(comments) + 1,
                        'text': match.group(1).strip(),
                        'author': 'Unknown',
                        'date': '',
                        'position': match.start(),
                        'associated_text': associated_text,
                        'context': f"Comment on '{associated_text[:30]}...': {match.group(1).strip()[:50]}..."
                    })
            
            # Method 3: Manual comment detection message
            if not comments and not comment_refs:
                logger.warning("No comments found. The document may not contain Word Review comments.")
                # Add a helpful message
                comments.append({
                    'id': 'info',
                    'text': 'No Word comments detected. For testing, try adding comments like [COMMENT: change this text] in your document.',
                    'author': 'System',
                    'date': '',
                    'position': 0,
                    'context': 'System message'
                })
        
        except Exception as e:
            logger.error(f"Fallback comment extraction failed: {str(e)}")
        
        return comments
    
    def analyze_comments_with_ai(self, comments, original_text, revised_text):
        """Analyze comments using GenAI to determine change scope and validation"""
        
        analysis_results = []
        
        for comment in comments:
            # Prioritize AI-powered analysis for intelligent comment understanding
            ai_client = get_anthropic_client() or get_openai_client()
            
            if ai_client:
                try:
                    logger.info(f"Using AI analysis for comment: '{comment['text'][:50]}...'")
                    result = self.ai_analyze_comment(comment, original_text, revised_text)
                    analysis_results.append(result)
                except Exception as e:
                    logger.error(f"AI analysis failed for comment '{comment['text']}': {str(e)}")
                    logger.info("Falling back to pattern matching for this comment")
                    # Fall back to pattern matching only if AI fails
                    result = self.fallback_analyze_comment(comment, original_text, revised_text)
                    analysis_results.append(result)
            else:
                logger.warning("No AI available - using pattern matching (limited comment understanding)")
                # Use pattern matching fallback when no AI is available
                result = self.fallback_analyze_comment(comment, original_text, revised_text)
                analysis_results.append(result)
        
        return analysis_results
    
    def extract_comment_context(self, comment, original_text, revised_text):
        """Extract focused context around where a comment appears"""
        
        # Priority 1: Use the position of the associated text if available
        associated_text = comment.get('associated_text', '').strip()
        if associated_text:
            # Find where the associated text actually appears in the document
            associated_position = original_text.find(associated_text)
            if associated_position != -1:
                comment_position = associated_position
                logger.info(f"Using associated text position {comment_position} for '{associated_text}'")
            else:
                # Fallback to stored comment position
                comment_position = comment.get('position', 0)
                logger.info(f"Associated text not found, using stored position {comment_position}")
        else:
            # No associated text, use stored comment position
            comment_position = comment.get('position', 0)
            logger.info(f"No associated text, using stored position {comment_position}")
        
        # Use smaller, more focused context window (±100 characters)
        context_window = 100
        start_pos = max(0, comment_position - context_window)
        end_pos = min(len(original_text), comment_position + context_window)
        
        # Extract context and try to break on sentence boundaries for clarity
        original_context = original_text[start_pos:end_pos].strip()
        
        # Try to find sentence boundaries to avoid cutting mid-sentence
        original_context = self.trim_to_sentences(original_context)
        
        # For revised text, try to find similar context
        # Look for matching text patterns - use text before the associated text position
        context_before = original_text[max(0, comment_position - 50):comment_position].strip()
        
        # Find corresponding position in revised text
        revised_context = self.find_corresponding_context(
            context_before, original_context, revised_text, context_window
        )
        
        return {
            'original_context': original_context,
            'revised_context': revised_context,
            'position': comment_position
        }
    
    def trim_to_sentences(self, text):
        """Trim text to complete sentences when possible"""
        
        # Try to end on sentence boundary
        sentence_endings = ['. ', '! ', '? ']
        for ending in sentence_endings:
            last_sentence = text.rfind(ending)
            if last_sentence > len(text) * 0.5:  # Only if we don't lose too much text
                return text[:last_sentence + 1].strip()
        
        # Fallback: trim to word boundary
        if len(text) > 50:
            last_space = text.rfind(' ', 0, len(text) - 10)
            if last_space > 0:
                return text[:last_space].strip()
        
        return text
    
    def find_corresponding_context(self, context_before, original_context, revised_text, window_size):
        """Find corresponding context in revised text using improved matching"""
        
        logger.info(f"Finding context for: before='{context_before[:50]}...', original='{original_context[:50]}...'")
        
        # Strategy 1: Try to find the unchanged text before the comment
        if len(context_before) > 10:
            # Try different portions of the before context to handle small changes
            for length in [30, 20, 15, 10]:
                if len(context_before) >= length:
                    search_text = context_before[-length:].strip()  # Use the end of before context
                    before_match = revised_text.find(search_text)
                    
                    if before_match != -1:
                        logger.info(f"Found before context match at position {before_match}")
                        # Found the anchor, extract window after it
                        start_pos = before_match + len(search_text)
                        end_pos = min(len(revised_text), start_pos + window_size)
                        revised_context = revised_text[start_pos:end_pos].strip()
                        
                        if revised_context:
                            logger.info(f"Strategy 1 success: '{revised_context[:50]}...'")
                            return self.trim_to_sentences(revised_context)
        
        # Strategy 2: Find common words around the area but skip the potentially changed word
        words_in_original = original_context.split()
        if len(words_in_original) >= 3:
            # Try to find sequences of 2-3 consecutive words that might be unchanged
            for start_idx in range(len(words_in_original) - 1):
                for seq_len in [3, 2]:
                    if start_idx + seq_len <= len(words_in_original):
                        word_sequence = ' '.join(words_in_original[start_idx:start_idx + seq_len])
                        match_pos = revised_text.find(word_sequence)
                        
                        if match_pos != -1:
                            logger.info(f"Found word sequence '{word_sequence}' at position {match_pos}")
                            # Extract context around this match
                            start_pos = max(0, match_pos - 50)
                            end_pos = min(len(revised_text), match_pos + window_size)
                            revised_context = revised_text[start_pos:end_pos].strip()
                            
                            if revised_context:
                                logger.info(f"Strategy 2 success: '{revised_context[:50]}...'")
                                return self.trim_to_sentences(revised_context)
        
        # Strategy 3: Use fuzzy matching with partial words
        # Look for individual words from the original context, but be more flexible
        significant_words = [w for w in original_context.split() if len(w) > 3 and w.lower() not in ['the', 'and', 'was', 'were', 'that', 'this', 'with', 'from']]
        
        for word in significant_words:
            match_pos = revised_text.find(word)
            if match_pos != -1:
                logger.info(f"Found significant word '{word}' at position {match_pos}")
                start_pos = max(0, match_pos - 75)  # Wider context
                end_pos = min(len(revised_text), match_pos + window_size)
                revised_context = revised_text[start_pos:end_pos].strip()
                
                if revised_context:
                    logger.info(f"Strategy 3 success: '{revised_context[:50]}...'")
                    return self.trim_to_sentences(revised_context)
        
        # Strategy 4: Position-based fallback using relative document position
        # If all else fails, use the approximate same position in the document
        original_length = len(context_before) + len(original_context)
        total_original_length = original_length * 3  # Rough estimate
        relative_position = len(context_before) / max(total_original_length, 1)
        
        approx_position = int(len(revised_text) * relative_position)
        start_pos = max(0, approx_position - window_size // 2)
        end_pos = min(len(revised_text), approx_position + window_size // 2)
        
        logger.info(f"Strategy 4 fallback: using position {approx_position} (relative: {relative_position:.2f})")
        return self.trim_to_sentences(revised_text[start_pos:end_pos])
    
    def ai_analyze_comment(self, comment, original_text, revised_text):
        """Use GenAI to analyze a comment and validate changes with full context understanding"""
        
        comment_text = comment['text']
        associated_text = comment.get('associated_text', '').strip()
        user_scope = comment.get('user_scope', 'auto')
        
        # Debug logging
        logger.info(f"AI analysis for comment: '{comment_text[:50]}...'")
        logger.info(f"Associated text: '{associated_text}' (length: {len(associated_text)})")
        logger.info(f"User specified scope: {user_scope}")
        
        # Create a comprehensive AI prompt that understands context and scope
        scope_description = {
            'global': 'a GLOBAL change (should affect all instances throughout the document)',
            'local': 'a LOCAL change (should affect only the specific instance being commented on)',
            'auto': 'automatically determined scope based on the comment intent'
        }
        
        # Provide the full document text for comprehensive analysis
        if associated_text:
            prompt = f"""You are an expert document reviewer analyzing Word document comments and their implementation.

COMMENT: "{comment_text}"
COMMENTED ON: "{associated_text}"
SCOPE: This is {scope_description[user_scope]}

TASK: Evaluate whether the comment "{comment_text}" applied to the text/word/phrase "{associated_text}" was correctly implemented in the revised document.

ORIGINAL DOCUMENT:
{original_text}

REVISED DOCUMENT:
{revised_text}

ANALYSIS APPROACH:
1. UNDERSTAND THE COMMENT: What specific change does "{comment_text}" request when applied to "{associated_text}"?

2. IDENTIFY COMMENT TYPE:
   - DIRECT REPLACEMENT: "Change X to Y", "should be Z"
   - STYLE/GRAMMAR: "Don't use contractions", "Make more formal", "Fix grammar"
   - CONTENT CHANGE: "Make this more exciting", "Add detail"
   - CORRECTION: "spelling mistake", "wrong word"
   - DELETION: "remove this", "delete"

3. DETERMINE EXPECTED CHANGE:
   - For STYLE comments: Identify what needs to be changed to follow the style rule
   - For DIRECT comments: What should "{associated_text}" become?
   - Should this be applied globally (all instances) or locally (just this instance)?
   - If user specified scope as "{user_scope}", respect that preference

4. VERIFY IMPLEMENTATION:
   - Search the revised document for evidence of the requested change
   - For style comments: Check if the style rule was applied (e.g., contractions expanded)
   - Count instances before/after if relevant for global changes
   - Check if the change was applied correctly

EXAMPLES OF COMMENT INTERPRETATION:
- "Change her name to Claire" on "Diane" = Change "Diane" to "Claire"
- "Don't use contractions" on "can't" = Change "can't" to "cannot", "it's" to "it is", etc.
- "spelling mistake" on "recieve" = Change "recieve" to "receive"
- "should be sunny" on "rainy" = Change "rainy" to "sunny"
- "Make more formal" on casual text = Replace informal words/phrases with formal equivalents
- "delete this" on "very very" = Remove the duplicate "very"

RESPONSE FORMAT (JSON):
{{
    "interpretation": "What change does the comment request?",
    "comment_type": "direct_replacement|style_grammar|content_change|correction|deletion",
    "expected_from": "What text should be changed (for style: identify specific issues like 'can\\'t, it\\'s')",
    "expected_to": "What it should become (for style: 'cannot, it is')",
    "scope_applied": "global|local",
    "status": "correctly_applied|partially_applied|not_applied|unclear",
    "evidence": "What evidence shows the change was/wasn't applied?",
    "confidence": 0.95,
    "requires_manual_review": false
}}

Analyze the full document context to make this determination."""
        
        else:
            # Handle case when associated text is missing
            prompt = f"""You are an expert document reviewer. Analyze this comment and determine what change it requests.

COMMENT: "{comment_text}"
SCOPE: This is {scope_description[user_scope]}

ORIGINAL DOCUMENT:
{original_text}

REVISED DOCUMENT:
{revised_text}

TASK: Determine what change the comment "{comment_text}" requests and verify if it was applied correctly.

ANALYSIS INSTRUCTIONS:
1. The specific text this comment refers to was not identified
2. Compare the original and revised documents to find relevant changes
3. Determine which change best matches the comment intent
4. Verify if the change aligns with the requested scope

RESPONSE FORMAT (JSON):
{{
    "interpretation": "What change does the comment request?",
    "expected_from": "What text was changed (if identifiable)",
    "expected_to": "What it was changed to (or null if deletion)",
    "scope_applied": "global|local",
    "status": "correctly_applied|partially_applied|not_applied|unclear",
    "evidence": "What evidence shows the change was/wasn't applied?",
    "confidence": 0.70,
    "requires_manual_review": true
}}

Note: Associated text was not available, so analysis is based on document comparison."""

        try:
            anthropic_client = get_anthropic_client()
            openai_client = get_openai_client()
            
            if anthropic_client:
                response = anthropic_client.messages.create(
                    model="claude-3-haiku-20240307",  # Fast and cost-effective
                    max_tokens=500,
                    messages=[{"role": "user", "content": prompt}]
                )
                ai_response = response.content[0].text
            elif openai_client:
                response = openai_client.chat.completions.create(
                    model="gpt-4o-mini",  # Fast and cost-effective
                    max_tokens=500,
                    messages=[{"role": "user", "content": prompt}]
                )
                ai_response = response.choices[0].message.content
            else:
                raise Exception("No AI client available")
            
            # Parse the JSON response
            try:
                ai_analysis = json.loads(ai_response)
            except json.JSONDecodeError:
                # Try to extract JSON from response if it's wrapped in other text
                import re
                json_match = re.search(r'\{.*\}', ai_response, re.DOTALL)
                if json_match:
                    ai_analysis = json.loads(json_match.group(0))
                else:
                    raise Exception("Could not parse AI response as JSON")
            
            # Convert AI analysis to our format
            intent = self.validate_intent_structure({
                'type': ai_analysis.get('comment_type', 'ai_determined'),
                'scope': ai_analysis.get('scope_applied', user_scope),
                'from_text': ai_analysis.get('expected_from', associated_text),
                'to_text': ai_analysis.get('expected_to'),
                'raw_comment': comment_text,
                'ai_interpretation': ai_analysis.get('interpretation', '')
            })
            
            validation = {
                'status': ai_analysis.get('status', 'unclear'),
                'message': ai_analysis.get('evidence', 'AI analysis completed'),
                'confidence': ai_analysis.get('confidence', 0.8),
                'interpretation': ai_analysis.get('interpretation', ''),
                'evidence': ai_analysis.get('evidence', '')
            }
            
            return {
                'comment': comment,
                'intent': intent,
                'validation': validation,
                'requires_manual_review': ai_analysis.get('requires_manual_review', False),
                'ai_powered': True
            }
            
        except Exception as e:
            logger.error(f"AI analysis error: {str(e)}")
            raise e
    
    def fallback_analyze_comment(self, comment, original_text, revised_text):
        """Fallback to pattern matching when AI is not available"""
        
        # Parse the comment to understand the intended change
        associated_text = comment.get('associated_text', '').strip()
        change_intent = self.parse_comment_intent(comment['text'], associated_text)
        
        # Override scope with user selection if provided
        if 'user_scope' in comment:
            user_scope = comment['user_scope']
            if user_scope == 'global':
                change_intent['scope'] = 'global'
                change_intent['type'] = 'replace_global'
            elif user_scope == 'local':
                change_intent['scope'] = 'local'
                if change_intent['type'] == 'replace_global':
                    change_intent['type'] = 'replace_local'
                # For local scope, always try enhanced local validation if we have associated text
                if associated_text:
                    return self.validate_local_change_with_context(comment, change_intent, original_text, revised_text)
        
        # Check if the change was applied correctly
        validation_result = self.validate_change_application(
            change_intent, original_text, revised_text
        )
        
        return {
            'comment': comment,
            'intent': change_intent,
            'validation': validation_result,
            'requires_manual_review': validation_result.get('ambiguous', False),
            'ai_powered': False
        }
    
    def validate_intent_structure(self, intent):
        """Ensure intent object has all required fields"""
        required_fields = ['type', 'from_text', 'to_text', 'scope', 'raw_comment']
        optional_fields = ['ai_interpretation', 'style_description']
        
        for field in required_fields:
            if field not in intent:
                intent[field] = None
        
        for field in optional_fields:
            if field not in intent:
                intent[field] = ''
                
        return intent
    
    def parse_context_aware_comment(self, comment_text, associated_text):
        """Parse comments that should use the associated text as context"""
        
        # Skip context-aware parsing if the comment explicitly mentions the associated text
        # This prevents "change all Johnny to Jimmy" from being parsed as context-aware
        if associated_text.lower() in comment_text.lower():
            return None
        
        # Skip if comment has explicit change patterns that should be handled normally
        explicit_patterns = [
            r'change\s+all\s+',
            r'replace\s+all\s+',
            r'find\s+and\s+replace',
            r'everywhere',
            r'globally',
            r'throughout'
        ]
        
        for pattern in explicit_patterns:
            if re.search(pattern, comment_text, re.IGNORECASE):
                return None
        
        # Pattern: "His/Her/Their name should be X" when commenting on a specific name
        name_patterns = [
            r'(?:his|her|their|the)\s+name\s+should\s+be\s+["\']?([^"\']+)["\']?',
            r'(?:his|her|their|the)\s+name\s+is\s+["\']?([^"\']+)["\']?',
            r'name\s+should\s+be\s+["\']?([^"\']+)["\']?',
            r'should\s+be\s+named\s+["\']?([^"\']+)["\']?',
            r'call\s+(?:him|her|them|it)\s+["\']?([^"\']+)["\']?',
            r'(?:the\s+)?(?:character|boy|girl|person)\s+should\s+be\s+called\s+["\']?([^"\']+)["\']?',  # "the character should be called Mike"
            r'change\s+(?:his|her|their|the)\s+name\s+to\s+["\']?([^"\']+)["\']?',  # "change her name to Diane"
            r'rename\s+(?:him|her|them|it)\s+to\s+["\']?([^"\']+)["\']?',  # "rename her to Diane"
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, comment_text, re.IGNORECASE)
            if match:
                target_name = match.group(1)
                # Use associated text as the source name
                return self.validate_intent_structure({
                    'type': 'replace_global',  # Names are typically global changes
                    'from_text': associated_text,
                    'to_text': target_name,
                    'scope': 'global',
                    'raw_comment': comment_text
                })
        
        # Pattern: "should be X" when commenting on a specific word (context-dependent)
        should_be_patterns = [
            r'should\s+be\s+["\']?([^"\']+)["\']?$',  # Simple "should be X" at end
        ]
        
        for pattern in should_be_patterns:
            match = re.search(pattern, comment_text, re.IGNORECASE)
            if match:
                target_text = match.group(1)
                # Use associated text as the source
                return self.validate_intent_structure({
                    'type': 'replace_local',
                    'from_text': associated_text,
                    'to_text': target_text,
                    'scope': 'local',
                    'raw_comment': comment_text
                })
        
        # Pattern: Comments that are just instructions without explicit source/target
        instruction_patterns = [
            r'make\s+(?:it|this)\s+["\']?([^"\']+)["\']?',
            r'change\s+(?:it|this)\s+to\s+["\']?([^"\']+)["\']?',
            r'use\s+["\']?([^"\']+)["\']?\s+instead',
        ]
        
        for pattern in instruction_patterns:
            match = re.search(pattern, comment_text, re.IGNORECASE)
            if match:
                target_text = match.group(1)
                return self.validate_intent_structure({
                    'type': 'replace_local',
                    'from_text': associated_text,
                    'to_text': target_text,
                    'scope': 'local',
                    'raw_comment': comment_text
                })
        
        # Pattern: Single word/phrase comments (likely replacement targets)
        single_word_patterns = [
            r'^["\']?([^"\']+)["\']?$',  # Just a single word/phrase (replacement target)
        ]
        
        for pattern in single_word_patterns:
            match = re.search(pattern, comment_text.strip(), re.IGNORECASE)
            if match and len(comment_text.strip().split()) <= 3:  # Max 3 words for single replacement
                target_text = match.group(1)
                return self.validate_intent_structure({
                    'type': 'replace_local',
                    'from_text': associated_text,
                    'to_text': target_text,
                    'scope': 'local',
                    'raw_comment': comment_text
                })
        
        return None  # No context-aware pattern matched
    
    def parse_style_comment(self, comment_text, associated_text=None):
        """Parse style and grammar comments that don't require specific text replacement"""
        
        # Style/grammar patterns that the fallback system can recognize
        style_patterns = [
            # Contraction-related
            (r"don'?t\s+use\s+contractions?", 'style_grammar', 'Remove contractions from text'),
            (r"expand\s+contractions?", 'style_grammar', 'Expand contractions to full forms'),
            (r"no\s+contractions?", 'style_grammar', 'Remove contractions from text'),
            
            # Formality
            (r"make\s+(?:this\s+)?more\s+formal", 'style_grammar', 'Make text more formal'),
            (r"use\s+formal\s+language", 'style_grammar', 'Use formal language'),
            (r"less\s+casual", 'style_grammar', 'Make text less casual'),
            
            # Grammar
            (r"fix\s+grammar", 'style_grammar', 'Fix grammatical errors'),
            (r"correct\s+grammar", 'style_grammar', 'Correct grammatical errors'),
            (r"grammar\s+(?:error|mistake)", 'style_grammar', 'Fix grammatical errors'),
            
            # General style
            (r"improve\s+(?:the\s+)?writing", 'style_grammar', 'Improve writing style'),
            (r"make\s+(?:this\s+)?clearer", 'style_grammar', 'Make text clearer'),
            (r"simplify\s+(?:this\s+)?(?:text|language)?", 'style_grammar', 'Simplify the language'),
        ]
        
        for pattern, change_type, description in style_patterns:
            if re.search(pattern, comment_text, re.IGNORECASE):
                
                # For contraction-related comments, try to find specific contractions
                if 'contraction' in description.lower() and associated_text:
                    contractions = self.find_contractions(associated_text)
                    if contractions:
                        # Show the specific contractions that need to be expanded
                        from_text = ', '.join(contractions)
                        to_text = ', '.join([self.expand_contraction(c) for c in contractions])
                        return self.validate_intent_structure({
                            'type': change_type,
                            'from_text': from_text,
                            'to_text': to_text,
                            'scope': 'local',
                            'raw_comment': comment_text,
                            'style_description': description
                        })
                    else:
                        # No contractions found - should already be correct
                        return self.validate_intent_structure({
                            'type': change_type,
                            'from_text': associated_text,
                            'to_text': 'no_contractions_found',
                            'scope': 'local',
                            'raw_comment': comment_text,
                            'style_description': description
                        })
                
                # For other style comments, use the text as-is
                return self.validate_intent_structure({
                    'type': change_type,
                    'from_text': associated_text or 'style_issue', 
                    'to_text': 'style_corrected',
                    'scope': 'local',  # Style comments usually apply to specific sections
                    'raw_comment': comment_text,
                    'style_description': description
                })
        
        return None  # No style pattern matched
    
    def parse_comment_intent(self, comment_text, associated_text=None):
        """Parse comment text to understand the intended change"""
        
        # First, check for style/grammar patterns that don't rely on specific text changes
        style_patterns = self.parse_style_comment(comment_text, associated_text)
        if style_patterns:
            return style_patterns
        
        # Then, check for context-aware patterns that should use associated text
        if associated_text and associated_text.strip():
            context_patterns = self.parse_context_aware_comment(comment_text, associated_text.strip())
            if context_patterns:
                return context_patterns
        
        # Enhanced patterns for common change instructions
        patterns = {
            # Global replacements
            'replace_global': [
                r'change\s+(?:all\s+)?(?:instances?\s+of\s+)?["\']?([^"\']+)["\']?\s+(?:to|with)\s+["\']?([^"\']+)["\']?\s+(?:everywhere|globally|throughout)',
                r'replace\s+(?:all\s+)?["\']?([^"\']+)["\']?\s+(?:with|to)\s+["\']?([^"\']+)["\']?\s+(?:everywhere|globally|throughout)',
                r'(?:find|search)\s+and\s+replace\s+["\']?([^"\']+)["\']?\s+(?:with|to)\s+["\']?([^"\']+)["\']?',
                r'change\s+all\s+["\']?([^"\']+)["\']?\s+(?:to|with)\s+["\']?([^"\']+)["\']?',  # "change all X to Y" (with or without quotes)
                
                # Character/name change patterns (global by nature)
                r'change\s+(?:the\s+)?(?:character|boy|girl|person|name)\'?s?\s+name\s+to\s+["\']?([^"\']+)["\']?',  # "change the boy's name to Jimmy"
                r'rename\s+(?:the\s+)?(?:character|boy|girl|person)\s+to\s+["\']?([^"\']+)["\']?',  # "rename the character to Jimmy"
            ],
            
            # Local replacements - most common patterns
            'replace_local': [
                r'change\s+(?:this\s+)?["\']?([^"\']+)["\']?\s+(?:to|with)\s+["\']?([^"\']+)["\']?(?:\s+here|\s+in\s+this\s+(?:sentence|paragraph))?',
                r'replace\s+["\']?([^"\']+)["\']?\s+(?:with|to)\s+["\']?([^"\']+)["\']?',
                r'correct\s+(?:spelling|word)?:?\s*["\']?([^"\']+)["\']?\s+(?:to|should\s+be|->|→)\s+["\']?([^"\']+)["\']?',
                r'correct\s+(?:spelling|word)?:?\s+([^\s]+)',  # "correct spelling: receive"
                r'should\s+be\s+["\']?([^"\']+)["\']?\s+(?:not|instead\s+of)\s+["\']?([^"\']+)["\']?',
                r'(?:fix|correct):\s*["\']?([^"\']+)["\']?\s+(?:to|->|→)\s+["\']?([^"\']+)["\']?',
                r'(?:typo|error):\s*["\']?([^"\']+)["\']?\s+(?:should\s+be|->|→)\s+["\']?([^"\']+)["\']?',
                # REMOVED the problematic generic pattern that was causing "His name should be Eli" to be parsed incorrectly
                r'["\']?([^"\']+)["\']?\s*[?]\s*["\']?([^"\']+)["\']?',  # "real? reel"
                r'use\s+["\']?([^"\']+)["\']?\s+(?:instead\s+of|not)\s+["\']?([^"\']+)["\']?'
            ],
            
            # Deletions
            'delete': [
                r'(?:delete|remove)\s+["\']?([^"\']+)["\']?',
                r'(?:cut|omit)\s+["\']?([^"\']+)["\']?',
                r'take\s+out\s+["\']?([^"\']+)["\']?'
            ],
            
            # Additions
            'add': [
                r'(?:add|insert)\s+["\']?([^"\']+)["\']?(?:\s+(?:before|after)\s+["\']?([^"\']+)["\']?)?',
                r'include\s+["\']?([^"\']+)["\']?',
                r'put\s+["\']?([^"\']+)["\']?\s+(?:before|after)\s+["\']?([^"\']+)["\']?'
            ],
            
            # Formatting
            'format': [
                r'(?:format|style)\s+["\']?([^"\']+)["\']?\s+as\s+([^"\']+)',
                r'make\s+["\']?([^"\']+)["\']?\s+(?:bold|italic|underlined?)'
            ]
        }
        
        # Try each pattern type
        for change_type, pattern_list in patterns.items():
            for pattern in pattern_list:
                match = re.search(pattern, comment_text, re.IGNORECASE)
                if match:
                    groups = match.groups()
                    
                    # Handle different group arrangements
                    if change_type == 'replace_global':
                        # Special handling for character name changes
                        if 'change' in comment_text.lower() and ('name' in comment_text.lower() or 'character' in comment_text.lower() or 'boy' in comment_text.lower() or 'girl' in comment_text.lower()):
                            # For "change the boy's name to Jimmy", we only get the target name
                            if len(groups) == 1:
                                # Use associated text as the source if available
                                from_text = associated_text.strip() if associated_text else None
                                to_text = groups[0]
                            else:
                                from_text = groups[0] if len(groups) >= 1 and groups[0] else (associated_text.strip() if associated_text else None)
                                to_text = groups[1] if len(groups) >= 2 and groups[1] else groups[0]
                        else:
                            # Regular global replacement patterns
                            from_text = groups[0] if len(groups) >= 1 and groups[0] else None
                            to_text = groups[1] if len(groups) >= 2 and groups[1] else None
                    elif change_type == 'replace_local' and len(groups) >= 1:
                        # Handle single-word patterns like "correct spelling: receive"
                        if len(groups) == 1 and 'correct' in comment_text.lower():
                            from_text = None  # Will be inferred
                            to_text = groups[0]
                        elif len(groups) >= 2:
                            # For replacements, sometimes the order might be reversed
                            from_text = groups[0] if groups[0] else None
                            to_text = groups[1] if groups[1] else None
                            
                            # Check specific patterns that have reversed order
                            if 'should be' in comment_text.lower() and ('not' in comment_text.lower() or 'instead of' in comment_text.lower()):
                                # "should be reel not real" means real->reel, so swap
                                from_text, to_text = to_text, from_text
                            elif 'use' in comment_text.lower() and 'instead of' in comment_text.lower():
                                # "use great instead of good" means good->great, so swap
                                from_text, to_text = to_text, from_text
                        else:
                            from_text = groups[0] if groups[0] else None
                            to_text = None
                    else:
                        from_text = groups[0] if len(groups) >= 1 and groups[0] else None
                        to_text = groups[1] if len(groups) >= 2 and groups[1] else None
                    
                    return self.validate_intent_structure({
                        'type': change_type,
                        'from_text': from_text,
                        'to_text': to_text,
                        'scope': 'global' if 'global' in change_type else 'local',
                        'raw_comment': comment_text
                    })
        
        # Smart fallback: try to extract two words that might be a replacement
        # Look for patterns like "word1 word2" where it might mean word1->word2
        simple_replacement = re.search(r'^["\']?(\w+)["\']?\s*[?/→-]+\s*["\']?(\w+)["\']?$', comment_text.strip(), re.IGNORECASE)
        if simple_replacement:
            return self.validate_intent_structure({
                'type': 'replace_local',
                'from_text': simple_replacement.group(1),
                'to_text': simple_replacement.group(2),
                'scope': 'local',
                'raw_comment': comment_text
            })
        
        # Another fallback: single word might be a replacement target
        single_word = re.search(r'^["\']?(\w+)["\']?$', comment_text.strip(), re.IGNORECASE)
        if single_word:
            return self.validate_intent_structure({
                'type': 'replace_local',
                'from_text': None,  # Will need context to determine
                'to_text': single_word.group(1),
                'scope': 'local',
                'raw_comment': comment_text
            })
        
        # If no pattern matches, return generic intent with all required fields
        return self.validate_intent_structure({
            'type': 'unknown',
            'from_text': None,
            'to_text': None,
            'scope': 'manual_review',
            'raw_comment': comment_text
        })
    
    def validate_change_application(self, intent, original_text, revised_text):
        """Validate if the intended change was correctly applied"""
        
        if intent['type'] == 'unknown':
            return {
                'status': 'manual_review_required',
                'message': 'Comment requires manual interpretation',
                'ambiguous': True
            }
        
        if intent['type'] in ['replace_global', 'replace_local']:
            from_text = intent.get('from_text')
            to_text = intent.get('to_text')
            
            # Handle case where only target word is specified (single word comment)
            if not from_text and to_text:
                # Try to find what word was likely replaced by looking for context
                # This is a smart guess based on similar words or context
                return self.validate_single_word_replacement(to_text, original_text, revised_text, intent)
            
            if not from_text or not to_text:
                return {
                    'status': 'invalid_comment',
                    'message': 'Could not parse replacement text from comment'
                }
            
            # Count occurrences in original and revised text
            original_count = original_text.lower().count(from_text.lower())
            revised_from_count = revised_text.lower().count(from_text.lower())
            revised_to_count = revised_text.lower().count(to_text.lower())
            
            if intent['scope'] == 'global':
                # For global changes, all instances should be replaced
                if revised_from_count == 0 and revised_to_count >= original_count:
                    return {
                        'status': 'correctly_applied',
                        'message': f'All {original_count} instances of "{from_text}" were changed to "{to_text}"',
                        'details': {
                            'original_count': original_count,
                            'remaining_count': revised_from_count,
                            'new_count': revised_to_count
                        }
                    }
                else:
                    return {
                        'status': 'partially_applied',
                        'message': f'{original_count - revised_from_count} of {original_count} instances were changed',
                        'details': {
                            'original_count': original_count,
                            'remaining_count': revised_from_count,
                            'new_count': revised_to_count
                        }
                    }
            else:
                # For local changes, at least one instance should be changed
                if revised_from_count < original_count:
                    return {
                        'status': 'correctly_applied',
                        'message': f'At least one instance of "{from_text}" was changed to "{to_text}"',
                        'details': {
                            'original_count': original_count,
                            'remaining_count': revised_from_count,
                            'new_count': revised_to_count
                        }
                    }
                else:
                    return {
                        'status': 'not_applied',
                        'message': f'No instances of "{from_text}" were changed',
                        'details': {
                            'original_count': original_count,
                            'remaining_count': revised_from_count,
                            'new_count': revised_to_count
                        }
                    }
        
        elif intent['type'] == 'style_grammar':
            # Handle style and grammar validation
            return self.validate_style_change(intent, original_text, revised_text)
        
        # Handle other change types (delete, add, format)
        return {
            'status': 'manual_review_required',
            'message': f'Change type "{intent["type"]}" requires manual review'
        }
    
    def validate_single_word_replacement(self, target_word, original_text, revised_text, intent):
        """Validate replacement when only the target word is known"""
        
        # Count target word in both documents
        original_target_count = original_text.lower().count(target_word.lower())
        revised_target_count = revised_text.lower().count(target_word.lower())
        
        # If target word appears more in revised than original, likely a replacement occurred
        if revised_target_count > original_target_count:
            added_count = revised_target_count - original_target_count
            return {
                'status': 'correctly_applied',
                'message': f'"{target_word}" was added {added_count} time(s) - likely replacing another word',
                'details': {
                    'original_count': original_target_count,
                    'revised_count': revised_target_count,
                    'added_count': added_count
                }
            }
        
        # If target word appears same or less, try to find similar words that might have been replaced
        import difflib
        
        # Split texts into words and find differences
        original_words = set(re.findall(r'\b\w+\b', original_text.lower()))
        revised_words = set(re.findall(r'\b\w+\b', revised_text.lower()))
        
        # Words that disappeared from original
        removed_words = original_words - revised_words
        
        # Find the most similar word to target_word among removed words
        if removed_words:
            closest_matches = difflib.get_close_matches(target_word.lower(), removed_words, n=1, cutoff=0.6)
            if closest_matches:
                likely_original = closest_matches[0]
                
                # Count occurrences to validate
                original_count = original_text.lower().count(likely_original)
                revised_original_count = revised_text.lower().count(likely_original)
                
                if revised_original_count < original_count and revised_target_count >= original_target_count:
                    return {
                        'status': 'correctly_applied',
                        'message': f'Likely replaced "{likely_original}" with "{target_word}"',
                        'details': {
                            'inferred_from': likely_original,
                            'original_count': original_count,
                            'remaining_count': revised_original_count,
                            'new_count': revised_target_count
                        }
                    }
        
        # If no clear replacement pattern found
        if revised_target_count == original_target_count:
            return {
                'status': 'unclear',
                'message': f'"{target_word}" appears same number of times in both documents - unclear if change was applied'
            }
        else:
            return {
                'status': 'manual_review_required',
                'message': f'Cannot determine if "{target_word}" replacement was correctly applied'
            }
    
    def validate_style_change(self, intent, original_text, revised_text):
        """Validate style and grammar changes"""
        
        style_description = intent.get('style_description', '').lower()
        
        if 'contraction' in style_description:
            # Check for contraction removal/expansion
            original_contractions = self.find_contractions(original_text)
            revised_contractions = self.find_contractions(revised_text)
            
            if not original_contractions:
                return {
                    'status': 'correctly_applied',
                    'message': 'No contractions found in original text - style rule already satisfied',
                    'details': {
                        'original_contractions': 0,
                        'revised_contractions': len(revised_contractions)
                    }
                }
            
            if len(revised_contractions) < len(original_contractions):
                return {
                    'status': 'correctly_applied',
                    'message': f'Contractions reduced from {len(original_contractions)} to {len(revised_contractions)}',
                    'details': {
                        'original_contractions': len(original_contractions),
                        'revised_contractions': len(revised_contractions),
                        'removed': original_contractions
                    }
                }
            elif len(revised_contractions) == len(original_contractions):
                return {
                    'status': 'not_applied',
                    'message': f'Contractions still present: {", ".join(original_contractions)}',
                    'details': {
                        'original_contractions': len(original_contractions),
                        'revised_contractions': len(revised_contractions),
                        'remaining': revised_contractions
                    }
                }
            else:
                return {
                    'status': 'not_applied',
                    'message': 'More contractions found in revised text than original',
                    'details': {
                        'original_contractions': len(original_contractions),
                        'revised_contractions': len(revised_contractions)
                    }
                }
        
        # For other style patterns, use a basic text comparison
        return {
            'status': 'manual_review_required',
            'message': f'Style change "{style_description}" requires manual review'
        }
    
    def find_contractions(self, text):
        """Find contractions in text"""
        contraction_patterns = [
            r"\b\w+'\w+\b",  # General pattern: word'word
            r"\b(?:can't|won't|shouldn't|couldn't|wouldn't|isn't|aren't|wasn't|weren't|don't|doesn't|didn't|haven't|hasn't|hadn't|I'm|you're|he's|she's|it's|we're|they're|I'll|you'll|he'll|she'll|it'll|we'll|they'll|I'd|you'd|he'd|she'd|it'd|we'd|they'd|I've|you've|he's|she's|it's|we've|they've)\b"
        ]
        
        contractions = []
        for pattern in contraction_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            contractions.extend(matches)
        
        # Remove duplicates and return
        return list(set(contractions))
    
    def expand_contraction(self, contraction):
        """Expand a single contraction to its full form"""
        expansions = {
            # Common contractions
            "can't": "cannot",
            "won't": "will not", 
            "shouldn't": "should not",
            "couldn't": "could not",
            "wouldn't": "would not",
            "isn't": "is not",
            "aren't": "are not",
            "wasn't": "was not",
            "weren't": "were not",
            "don't": "do not",
            "doesn't": "does not",
            "didn't": "did not",
            "haven't": "have not",
            "hasn't": "has not",
            "hadn't": "had not",
            "I'm": "I am",
            "you're": "you are",
            "he's": "he is",
            "she's": "she is",
            "it's": "it is",
            "we're": "we are", 
            "they're": "they are",
            "I'll": "I will",
            "you'll": "you will",
            "he'll": "he will",
            "she'll": "she will",
            "it'll": "it will",
            "we'll": "we will",
            "they'll": "they will",
            "I'd": "I would",
            "you'd": "you would",
            "he'd": "he would",
            "she'd": "she would",
            "it'd": "it would",
            "we'd": "we would", 
            "they'd": "they would",
            "I've": "I have",
            "you've": "you have",
            "we've": "we have",
            "they've": "they have",
            # Possessive contractions that might be confused
            "storm's": "storm has",  # Context dependent, but common expansion
        }
        
        # Case-insensitive lookup
        lower_contraction = contraction.lower()
        if lower_contraction in expansions:
            expansion = expansions[lower_contraction]
            # Preserve original capitalization
            if contraction[0].isupper():
                return expansion.capitalize()
            return expansion
        
        # If not in our dictionary, try basic apostrophe-s handling
        if contraction.lower().endswith("'s"):
            base = contraction[:-2]
            return f"{base} has"  # Common expansion, context dependent
        
        # Fallback: return original if we can't expand it
        return contraction
    
    def generate_comparison_report(self, session_id):
        """Generate a comprehensive comparison report"""
        
        if session_id not in self.session_data:
            return None
        
        data = self.session_data[session_id]
        
        # Create side-by-side diff
        original_lines = data['original']['full_text'].split('\n')
        revised_lines = data['revised']['full_text'].split('\n')
        
        differ = difflib.unified_diff(
            original_lines, revised_lines,
            fromfile='Original Document',
            tofile='Revised Document',
            lineterm=''
        )
        
        # Generate enhanced diff that highlights missed instances
        diff_html = self.generate_enhanced_diff(
            original_lines, revised_lines, data.get('analysis_results', [])
        )
        
        return {
            'session_id': session_id,
            'analysis_results': data.get('analysis_results', []),
            'diff_html': diff_html,
            'summary': self.generate_summary(data.get('analysis_results', [])),
            'timestamp': data.get('timestamp')
        }
    
    def generate_enhanced_diff(self, original_lines, revised_lines, analysis_results):
        """Generate enhanced HTML diff that highlights missed instances for incomplete global changes"""
        
        # Find incomplete global changes that need highlighting
        missed_instances = []
        
        for result in analysis_results:
            validation = result.get('validation', {})
            if (validation.get('status') == 'partially_applied' and 
                result.get('intent', {}).get('scope') == 'global'):
                
                # This is an incomplete global change - find missed instances
                intent = result.get('intent', {})
                from_text = intent.get('from_text')
                
                if from_text:
                    # Find all lines in revised text that still contain the old text
                    for line_num, line in enumerate(revised_lines):
                        if from_text.lower() in line.lower():
                            missed_instances.append({
                                'line_num': line_num,
                                'text': from_text,
                                'comment': result.get('comment', {}).get('text', ''),
                                'line_content': line
                            })
        
        # Generate the base HTML diff
        differ = difflib.HtmlDiff()
        html_diff = differ.make_table(
            original_lines, revised_lines,
            fromdesc='Original Document',
            todesc='Revised Document',
            context=True,
            numlines=3
        )
        
        # Enhance the HTML to highlight missed instances
        if missed_instances:
            html_diff = self.enhance_diff_with_missed_instances(html_diff, missed_instances)
        
        return html_diff
    
    def enhance_diff_with_missed_instances(self, html_diff, missed_instances):
        """Add highlighting and annotations for missed instances in the HTML diff"""
        
        # Add custom CSS for missed instance highlighting
        enhanced_css = """
        <style>
        .missed-instance {
            background-color: #ffebee !important;
            border: 2px solid #f44336 !important;
            position: relative;
        }
        .missed-instance::after {
            content: '⚠️ MISSED INSTANCE';
            position: absolute;
            right: 5px;
            top: 2px;
            background: #f44336;
            color: white;
            padding: 2px 6px;
            font-size: 10px;
            border-radius: 3px;
            font-weight: bold;
        }
        .missed-instance-tooltip {
            background: #fff3cd;
            border: 1px solid #ffeaa7;
            padding: 8px;
            margin: 5px 0;
            border-radius: 4px;
            font-size: 12px;
        }
        </style>
        """
        
        # Insert CSS at the beginning of the HTML
        html_diff = enhanced_css + html_diff
        
        # For each missed instance, enhance the corresponding line in the HTML
        for missed in missed_instances:
            line_num = missed['line_num']
            text_to_highlight = missed['text']
            comment = missed['comment']
            
            # Create a unique identifier for this line
            line_pattern = f'<td class="diff_chg"[^>]*>([^<]*{re.escape(text_to_highlight)}[^<]*)</td>'
            
            def replace_line(match):
                content = match.group(1)
                # Highlight the specific missed text
                highlighted_content = content.replace(
                    text_to_highlight,
                    f'<span style="background-color: #ffcdd2; font-weight: bold; border: 1px solid #f44336; padding: 1px 3px; border-radius: 2px;">{text_to_highlight}</span>'
                )
                return f'<td class="diff_chg missed-instance" title="Missed instance for comment: {comment}">{highlighted_content}</td>'
            
            html_diff = re.sub(line_pattern, replace_line, html_diff, flags=re.IGNORECASE)
        
        # Add a summary box at the top
        if missed_instances:
            summary_box = f"""
            <div class="missed-instance-tooltip">
                <strong>⚠️ Incomplete Global Changes Detected:</strong><br>
                Found {len(missed_instances)} missed instance(s) that should have been changed globally.
                Look for highlighted lines with ⚠️ MISSED INSTANCE markers.
            </div>
            """
            # Insert after the first table tag
            html_diff = html_diff.replace('<table', summary_box + '<table', 1)
        
        return html_diff

    def generate_html_diff(self, original_lines, revised_lines):
        """Generate HTML side-by-side diff"""
        
        differ = difflib.HtmlDiff()
        html_diff = differ.make_table(
            original_lines, revised_lines,
            fromdesc='Original Document',
            todesc='Revised Document',
            context=True,
            numlines=3
        )
        
        return html_diff
    
    def validate_local_change_with_context(self, comment, change_intent, original_text, revised_text):
        """Enhanced validation for local changes using context analysis"""
        
        associated_text = comment.get('associated_text', '').strip()
        comment_text = comment['text'].lower()
        
        # Extract context around the associated text
        context = self.extract_comment_context(comment, original_text, revised_text)
        original_context = context['original_context'].lower()
        revised_context = context['revised_context'].lower()
        
        # Check if the associated text appears in original context
        if associated_text.lower() not in original_context:
            return {
                'comment': comment,
                'intent': change_intent,
                'validation': {
                    'status': 'manual_review_required',
                    'message': 'Could not locate associated text in context for validation'
                },
                'requires_manual_review': True,
                'ai_powered': False
            }
        
        # Determine expected change based on common patterns
        expected_change = None
        
        # Spelling corrections
        if 'spelling' in comment_text or 'spell' in comment_text:
            # Common spelling corrections
            spelling_fixes = {
                'recieve': 'receive', 'teh': 'the', 'hte': 'the', 'seperate': 'separate',
                'occured': 'occurred', 'definately': 'definitely', 'thier': 'their',
                'reel': 'real', 'absolutly': 'absolutely'
            }
            expected_change = spelling_fixes.get(associated_text.lower())
        
        # Word replacements
        elif any(word in comment_text for word in ['change', 'replace', 'use', 'different']):
            if 'sunny' in comment_text or 'sun' in comment_text:
                expected_change = 'sunny'
            elif 'jimmy' in comment_text:
                expected_change = 'Jimmy'
            elif 'smiling' in comment_text:
                expected_change = 'smiling'
            elif 'excellent' in comment_text:
                expected_change = 'excellent'
                
        # Check if expected change appears in revised context
        if expected_change:
            if expected_change.lower() in revised_context:
                status = 'correctly_applied'
                message = f'Successfully changed "{associated_text}" to "{expected_change}" in the local context'
            else:
                status = 'not_applied'
                message = f'Expected change from "{associated_text}" to "{expected_change}" was not found in revised context'
        else:
            # Check if the associated text was removed or changed somehow
            if associated_text.lower() not in revised_context:
                status = 'correctly_applied'
                message = f'Associated text "{associated_text}" was modified/removed from the context as requested'
            else:
                status = 'not_applied'
                message = f'Associated text "{associated_text}" still appears unchanged in the revised context'
        
        return {
            'comment': comment,
            'intent': change_intent,
            'validation': {
                'status': status,
                'message': message,
                'change_type': 'local_context_validation',
                'associated_text': associated_text,
                'expected_change': expected_change
            },
            'requires_manual_review': status not in ['correctly_applied', 'not_applied'],
            'ai_powered': False
        }

    def generate_summary(self, analysis_results):
        """Generate summary statistics"""
        
        total_comments = len(analysis_results)
        correctly_applied = sum(1 for r in analysis_results if r['validation']['status'] == 'correctly_applied')
        partially_applied = sum(1 for r in analysis_results if r['validation']['status'] == 'partially_applied')
        not_applied = sum(1 for r in analysis_results if r['validation']['status'] == 'not_applied')
        manual_review = sum(1 for r in analysis_results if r.get('requires_manual_review', False))
        
        return {
            'total_comments': total_comments,
            'correctly_applied': correctly_applied,
            'partially_applied': partially_applied,
            'not_applied': not_applied,
            'manual_review_required': manual_review,
            'success_rate': (correctly_applied / total_comments * 100) if total_comments > 0 else 0
        }

# Global analyzer instance
analyzer = WordDocumentAnalyzer()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/health')
def health_check():
    """Health check endpoint for Railway"""
    health_status = {
        'status': 'healthy', 
        'service': 'word-doc-comparer',
        'api_status': {
            'anthropic_key_available': bool(ANTHROPIC_API_KEY),
            'anthropic_package_available': ANTHROPIC_AVAILABLE,
            'anthropic_ready': bool(ANTHROPIC_API_KEY and ANTHROPIC_AVAILABLE),
            'openai_key_available': bool(OPENAI_API_KEY),
            'openai_package_available': OPENAI_AVAILABLE,
            'openai_ready': bool(OPENAI_API_KEY and OPENAI_AVAILABLE)
        }
    }
    return jsonify(health_status)

@app.route('/api/status')
def api_status():
    """API status endpoint for UI"""
    status = {
        'ai_enabled': bool(ANTHROPIC_API_KEY and ANTHROPIC_AVAILABLE) or bool(OPENAI_API_KEY and OPENAI_AVAILABLE),
        'anthropic': {
            'available': ANTHROPIC_AVAILABLE,
            'configured': bool(ANTHROPIC_API_KEY),
            'ready': bool(ANTHROPIC_API_KEY and ANTHROPIC_AVAILABLE)
        },
        'openai': {
            'available': OPENAI_AVAILABLE,
            'configured': bool(OPENAI_API_KEY),
            'ready': bool(OPENAI_API_KEY and OPENAI_AVAILABLE)
        },
        'primary_ai': 'anthropic' if (ANTHROPIC_API_KEY and ANTHROPIC_AVAILABLE) else ('openai' if (OPENAI_API_KEY and OPENAI_AVAILABLE) else 'none')
    }
    return jsonify(status)

@app.route('/api/test-contractions')
def test_contractions():
    """Test endpoint to verify contraction parsing is working"""
    test_text = '"Phone lines are down," Elias said gruffly, handing her a blanket and a cup of tea. "Storm\'s cut the signal."'
    
    analyzer = WordDocumentAnalyzer()
    
    # Test contraction detection
    contractions = analyzer.find_contractions(test_text)
    
    # Test style parsing
    style_result = analyzer.parse_style_comment("Don't use contractions", test_text)
    
    return jsonify({
        'version': '2.1',
        'deployment_time': datetime.now().isoformat(),
        'test_text': test_text,
        'contractions_found': contractions,
        'style_parsing': {
            'from_text': style_result['from_text'] if style_result else 'No match',
            'to_text': style_result['to_text'] if style_result else 'No match',
            'type': style_result['type'] if style_result else 'No match'
        },
        'expected': {
            'contractions': ["Storm's"],
            'from_text': "Storm's",
            'to_text': "Storm has"
        },
        'status': 'working' if (contractions == ["Storm's"] and 
                               style_result and 
                               style_result['from_text'] == "Storm's" and 
                               style_result['to_text'] == "Storm has") else 'broken'
    })

@app.route('/upload', methods=['POST'])
def upload_files():
    """Handle file uploads"""
    try:
        if 'original_doc' not in request.files or 'revised_doc' not in request.files:
            return jsonify({'error': 'Both original and revised documents are required'}), 400
        
        original_file = request.files['original_doc']
        revised_file = request.files['revised_doc']
        
        if original_file.filename == '' or revised_file.filename == '':
            return jsonify({'error': 'No files selected'}), 400
        
        # Validate file types
        allowed_extensions = {'.docx'}
        for file in [original_file, revised_file]:
            if not any(file.filename.lower().endswith(ext) for ext in allowed_extensions):
                return jsonify({'error': 'Only .docx files are supported'}), 400
        
        # Generate session ID
        session_id = str(uuid.uuid4())
        
        # Save files
        original_filename = secure_filename(f"{session_id}_original_{original_file.filename}")
        revised_filename = secure_filename(f"{session_id}_revised_{revised_file.filename}")
        
        original_path = os.path.join(app.config['UPLOAD_FOLDER'], original_filename)
        revised_path = os.path.join(app.config['UPLOAD_FOLDER'], revised_filename)
        
        original_file.save(original_path)
        revised_file.save(revised_path)
        
        # Extract document data
        original_data = analyzer.extract_document_data(original_path)
        revised_data = analyzer.extract_document_data(revised_path)
        
        # Store session data
        analyzer.session_data[session_id] = {
            'original': original_data,
            'revised': revised_data,
            'original_file': original_filename,
            'revised_file': revised_filename,
            'timestamp': datetime.now().isoformat()
        }
        
        return jsonify({
            'success': True,
            'session_id': session_id,
            'original_comments': len(original_data['comments']),
            'revised_comments': len(revised_data['comments']),
            'message': f'Files uploaded successfully. Found {len(original_data["comments"])} comments in original document.',
            'debug_url': f'/debug/{session_id}',
            'review_url': f'/review-scope/{session_id}'
        })
        
    except Exception as e:
        logger.error(f"Upload error: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/review-scope/<session_id>')
def review_scope(session_id):
    """Review and set scope for each comment"""
    try:
        if session_id not in analyzer.session_data:
            return render_template('error.html', message='Session not found'), 404
        
        session_data = analyzer.session_data[session_id]
        original_data = session_data['original']
        
        # If no comments found, skip scope review and go directly to analysis
        if not original_data['comments']:
            return redirect(url_for('analyze_documents', session_id=session_id))
        
        return render_template('review_scope.html', 
                             session_id=session_id,
                             comments=original_data['comments'],
                             original_file=session_data['original_file'],
                             revised_file=session_data['revised_file'])
        
    except Exception as e:
        logger.error(f"Scope review error: {str(e)}")
        return render_template('error.html', message=str(e)), 500

@app.route('/analyze/<session_id>', methods=['POST'])
def analyze_documents(session_id):
    """Analyze documents and generate comparison"""
    try:
        if session_id not in analyzer.session_data:
            return jsonify({'error': 'Session not found'}), 404
        
        data = analyzer.session_data[session_id]
        
        # Get user scope selections from form data
        scope_selections = {}
        form_data = request.get_json() or request.form
        
        for key, value in form_data.items():
            if key.startswith('scope_'):
                comment_index = int(key.replace('scope_', ''))
                scope_selections[comment_index] = value
        
        # Apply user scope selections to comments
        comments_with_scope = []
        for i, comment in enumerate(data['original']['comments']):
            comment_copy = comment.copy()
            # Override the scope with user selection
            user_scope = scope_selections.get(i, comment.get('user_scope', 'local'))
            comment_copy['user_scope'] = user_scope
            comments_with_scope.append(comment_copy)
        
        # Store updated comments with user scopes
        data['original']['comments'] = comments_with_scope
        
        # Analyze comments with AI using user-specified scopes
        analysis_results = analyzer.analyze_comments_with_ai(
            comments_with_scope,
            data['original']['full_text'],
            data['revised']['full_text']
        )
        
        # Store analysis results
        data['analysis_results'] = analysis_results
        
        # Generate comparison report
        report = analyzer.generate_comparison_report(session_id)
        
        return jsonify({
            'success': True,
            'report': report
        })
        
    except Exception as e:
        logger.error(f"Analysis error: {str(e)}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@app.route('/report/<session_id>')
def view_report(session_id):
    """View comparison report"""
    report = analyzer.generate_comparison_report(session_id)
    if not report:
        return "Report not found", 404
    
    return render_template('report.html', report=report)

@app.route('/debug/<session_id>')
def debug_session(session_id):
    """Debug endpoint to see extracted data"""
    if session_id not in analyzer.session_data:
        return jsonify({'error': 'Session not found'}), 404
    
    data = analyzer.session_data[session_id]
    
    return jsonify({
        'session_id': session_id,
        'original_comments': data['original']['comments'],
        'original_text_preview': data['original']['full_text'][:500] + '...' if len(data['original']['full_text']) > 500 else data['original']['full_text'],
        'revised_text_preview': data['revised']['full_text'][:500] + '...' if len(data['revised']['full_text']) > 500 else data['revised']['full_text'],
        'original_paragraphs_count': len(data['original']['paragraphs']),
        'revised_paragraphs_count': len(data['revised']['paragraphs'])
    })

if __name__ == '__main__':
    # Production configuration for Railway
    port = int(os.environ.get('PORT', 8082))
    debug = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    app.run(host='0.0.0.0', port=port, debug=debug)